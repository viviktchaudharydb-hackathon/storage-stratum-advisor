"""
ARB Decision Record generator - produces a formal Architecture Review Board
paper as a Word document (bytes, for st.download_button).

Two entry points:
  build_arb_document(state, fmt_money)            - single-domain record (v4, unchanged)
  build_blueprint_arb_document(...)               - consolidated record for a Solution
                                                    Blueprint: one document covering all
                                                    components + cross-domain view (v5)

Uses python-docx at runtime (Cloud Run). Sections follow a typical bank ARB
template: context, options considered, cost comparison (with the deterministic
model disclosed), procurement context, compliance results, recommendation,
and sign-off placeholders.
"""

import io
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1A, 0x47, 0x8F)
STATUS_COLORS = {"pass": RGBColor(0x1E, 0x7A, 0x33),
                 "warn": RGBColor(0xB0, 0x6A, 0x00),
                 "fail": RGBColor(0xB0, 0x1E, 0x1E)}
STATUS_LABEL = {"pass": "PASS", "warn": "REVIEW", "fail": "FAIL"}

# Worst-first ordering used to consolidate per-domain verdicts into one
_STATUS_RANK = {"fail": 0, "warn": 1, "pass": 2}


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _kv_table(doc, rows: List[tuple]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
    return table


def _small_italic(doc, text: str, size: int = 8):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.italic = True
    return p


def _status_run(paragraph, status: str, prefix: str = "", suffix: str = ""):
    """Append a colored PASS/REVIEW/FAIL run to a paragraph."""
    if prefix:
        paragraph.add_run(prefix)
    r = paragraph.add_run(STATUS_LABEL[status])
    r.bold = True
    r.font.color.rgb = STATUS_COLORS[status]
    if suffix:
        paragraph.add_run(suffix)
    return r


def _compliance_block(doc, compliance: List[Dict[str, Any]]):
    """Render per-vendor compliance results (shared by both document types)."""
    if not compliance:
        doc.add_paragraph("Compliance checks were not executed for this analysis.")
        return
    for block in compliance:
        p = doc.add_paragraph()
        p.add_run(f"{block['vendor']} - overall: ").bold = True
        _status_run(p, block["overall"])
        for c in block["checks"]:
            cp = doc.add_paragraph(style="List Bullet")
            sr = cp.add_run(f"[{STATUS_LABEL[c['status']]}] ")
            sr.bold = True
            sr.font.color.rgb = STATUS_COLORS[c["status"]]
            cp.add_run(f"{c['name']}: {c['detail']}")


def _tco_table(doc, tco_list: List[Dict[str, Any]], fmt_money, extra_first_col: str = None):
    """Render a TCO comparison table. If extra_first_col is given, an extra
    leading column (e.g. Domain) is added and each tco entry must carry that key."""
    cols = ["Vendor", "Model", "Negotiated Discount", "3-Yr TCO (est.)", "Range"]
    if extra_first_col:
        cols = [extra_first_col] + cols
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for j, h in enumerate(cols):
        hdr[j].text = h
        for run in hdr[j].paragraphs[0].runs:
            run.bold = True
    for t in tco_list:
        row = table.add_row().cells
        off = 0
        if extra_first_col:
            row[0].text = str(t.get("_extra", ""))
            off = 1
        row[off + 0].text = t["vendor"]
        row[off + 1].text = t["model"]
        row[off + 2].text = f"{t['negotiated_discount_pct']}%" if t["has_agreement"] else "-"
        row[off + 3].text = fmt_money(t["total_3yr"])
        row[off + 4].text = f"{fmt_money(t['range'][0])} – {fmt_money(t['range'][1])}"
    return table


def _methodology_note(doc):
    _small_italic(doc,
        "Methodology: deterministic model - domain list rate by cost profile (or cloud "
        "capacity tier), less negotiated discounts sourced from procurement agreements, "
        "plus facilities and migration baseline; ±15% sizing uncertainty band. "
        "Figures are pre-quotation estimates for comparison, not commercial offers.")


def _approvals(doc, roles: List[str]):
    _h(doc, "Approvals")
    table = doc.add_table(rows=len(roles) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for j, h in enumerate(["Role", "Name / Signature", "Date"]):
        hdr[j].text = h
        for run in hdr[j].paragraphs[0].runs:
            run.bold = True
    for i, role in enumerate(roles, start=1):
        table.rows[i].cells[0].text = role


def _disclaimer(doc):
    _small_italic(doc,
        "\nGenerated by the Enterprise Infrastructure Advisor. AI-assisted content "
        "(architecture analysis, vendor fit narrative) is marked by origin in the tool; "
        "cost figures and compliance results are computed deterministically. This record "
        "requires human review and approval before any procurement action.")


def _top_vendor(state: Dict[str, Any]):
    return next((r for r in state.get("vendor_recommendations", [])
                 if r.get("fit_score", 0) > 0), None)


def _component_verdict(state: Dict[str, Any]) -> str:
    """Consolidated verdict for one component = the top vendor's overall status,
    or worst-of-all if no top vendor is identifiable."""
    compliance = state.get("compliance_results", [])
    if not compliance:
        return "warn"
    top = _top_vendor(state)
    if top:
        for block in compliance:
            if block["vendor"] == top.get("name"):
                return block["overall"]
    return min((b["overall"] for b in compliance), key=lambda s: _STATUS_RANK[s])


# ---------------------------------------------------------------------------
# v4: single-domain ARB record (unchanged behavior)
# ---------------------------------------------------------------------------

def build_arb_document(state: Dict[str, Any], fmt_money) -> bytes:
    """state: the final AgentState dict. fmt_money: TCOEngine.fmt."""
    req = state["requirements"]
    report = state.get("final_report", {})
    arch = state.get("architecture_analysis", {})
    vendors = [v for v in state.get("vendor_recommendations", []) if v.get("fit_score", 0) > 0]
    tco_list = state.get("tco_estimates", [])
    tco_by_vendor = {t["vendor"]: t for t in tco_list}
    compliance = state.get("compliance_results", [])
    procurement = state.get("procurement_context", [])
    sources = state.get("market_data", {}).get("sources", [])

    doc = Document()

    # Title
    doc.add_heading("Architecture Review Board - Decision Record", level=0)
    sub = doc.add_paragraph(
        f"{req.domain} · {req.workload} · Generated {datetime.now():%d %b %Y %H:%M} "
        f"· Enterprise Infrastructure Advisor (Vertex AI)")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sub.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 1. Context
    _h(doc, "1. Context and Requirements")
    _kv_table(doc, [
        ("Domain", req.domain),
        ("Workload", req.workload),
        ("Scale", f"{req.capacity} {req.unit}"),
        ("Deployment model", req.deployment),
        ("Design priority", req.priority),
        ("Availability SLA", req.availability_target),
    ])

    # 2. Recommended architecture
    _h(doc, "2. Recommended Architecture")
    if arch.get("architecture_type"):
        p = doc.add_paragraph()
        p.add_run(f"Architecture pattern: {arch['architecture_type']}").bold = True
    for rec in arch.get("key_recommendations", []):
        doc.add_paragraph(rec, style="List Bullet")
    if arch.get("redundancy_approach"):
        doc.add_paragraph(f"Redundancy approach: {arch['redundancy_approach']}")
    if arch.get("scalability_notes"):
        doc.add_paragraph(f"Scalability: {arch['scalability_notes']}")

    # 3. Options considered
    _h(doc, "3. Options Considered")
    for i, v in enumerate(vendors[:4], 1):
        _h(doc, f"3.{i} {v['name']} - fit score {v.get('fit_score', 'N/A')}/10", level=2)
        if v.get("strengths"):
            doc.add_paragraph("Strengths:", style="Intense Quote")
            for s in v["strengths"]:
                doc.add_paragraph(s, style="List Bullet")
        if v.get("considerations"):
            doc.add_paragraph("Considerations:", style="Intense Quote")
            for c in v["considerations"]:
                doc.add_paragraph(c, style="List Bullet")

    # 4. Cost comparison
    _h(doc, "4. Three-Year TCO Comparison")
    if tco_list:
        _tco_table(doc, tco_list, fmt_money)
        _methodology_note(doc)

    # 5. Procurement context
    _h(doc, "5. Procurement Context (retrieved)")
    if procurement:
        for d in procurement:
            meta = d.get("meta", {})
            p = doc.add_paragraph()
            p.add_run(d["id"]).bold = True
            tags = []
            if meta.get("vendor") not in (None, "", "none", "multiple"):
                tags.append(f"vendor: {meta['vendor']}")
            if meta.get("discount_pct"):
                tags.append(f"discount: {meta['discount_pct']}%")
            if meta.get("valid_until"):
                tags.append(f"valid until: {meta['valid_until']}")
            if tags:
                p.add_run("  (" + " · ".join(tags) + ")").italic = True
            excerpt = d["text"][:350] + ("…" if len(d["text"]) > 350 else "")
            doc.add_paragraph(excerpt)
    else:
        doc.add_paragraph("No relevant internal procurement documents were retrieved for this analysis.")

    # 6. Compliance assessment
    _h(doc, "6. Compliance Assessment")
    _compliance_block(doc, compliance)

    # 7. Recommendation & next steps
    _h(doc, "7. Recommendation and Next Steps")
    if vendors:
        top = vendors[0]
        p = doc.add_paragraph()
        p.add_run(f"Recommended option: {top['name']}").bold = True
        tco = tco_by_vendor.get(top["name"])
        if tco:
            doc.add_paragraph(
                f"Estimated 3-year TCO {fmt_money(tco['range'][0])}–{fmt_money(tco['range'][1])}"
                + (" under the existing negotiated agreement." if tco["has_agreement"]
                   else "; no existing agreement - commercial negotiation required."))
    for step in report.get("next_steps", []):
        doc.add_paragraph(step, style="List Number")

    # 8. Market sources
    if sources:
        _h(doc, "8. Market Intelligence Sources")
        for s in sources:
            doc.add_paragraph(f"{s.get('title') or s.get('url')} - {s.get('url')}", style="List Bullet")

    _approvals(doc, ["Solution Architect", "Architecture Review Board Chair",
                     "Procurement / Vendor Management"])
    _disclaimer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# v5: consolidated blueprint-level ARB record
# ---------------------------------------------------------------------------

def build_blueprint_arb_document(blueprint_name: str,
                                 description: str,
                                 driver_label: str,
                                 driver_value: int,
                                 params: Dict[str, float],
                                 param_labels: Dict[str, str],
                                 component_results: List[Dict[str, Any]],
                                 synergy: Dict[str, Any],
                                 fmt_money) -> bytes:
    """One ARB decision record for an entire Solution Blueprint.

    Args:
        blueprint_name:    key from BLUEPRINTS (e.g. "Core Banking Modernization").
        description:       BLUEPRINTS[name]["description"].
        driver_label:      BLUEPRINTS[name]["driver"]["label"].
        driver_value:      the driver value the user selected.
        params:            resolved sizing params actually used (user-adjusted).
        param_labels:      {param_key: label} from BLUEPRINTS[name]["params"] for display.
        component_results: [{domain, workload, capacity, unit, rationale, state}] -
                           same list handed to analyze_synergy(), with each
                           component's final AgentState dict under "state".
        synergy:           output of blueprints.analyze_synergy(component_results).
        fmt_money:         TCOEngine.fmt.
    """
    doc = Document()

    # Title
    doc.add_heading("Architecture Review Board - Solution Decision Record", level=0)
    sub = doc.add_paragraph(
        f"Solution Blueprint: {blueprint_name} · Generated {datetime.now():%d %b %Y %H:%M} "
        f"· Enterprise Infrastructure Advisor (Vertex AI)")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sub.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Consolidated verdict: worst of the per-component verdicts
    verdicts = {c["domain"]: _component_verdict(c["state"]) for c in component_results}
    overall = min(verdicts.values(), key=lambda s: _STATUS_RANK[s]) if verdicts else "warn"

    # 1. Solution context
    _h(doc, "1. Solution Context")
    doc.add_paragraph(description)
    rows = [(driver_label, driver_value)]
    rows += [(param_labels.get(k, k), v) for k, v in params.items()]
    _kv_table(doc, rows)
    _small_italic(doc,
        "Component capacities below are derived deterministically from the driver metric "
        "and the sizing assumptions above (formulas in blueprints.py) - the domains are "
        "correlated, not analyzed in isolation.")

    # 2. Component sizing
    _h(doc, "2. Component Sizing (derived)")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for j, h in enumerate(["Domain", "Workload", "Derived Capacity", "Sizing Rationale"]):
        hdr[j].text = h
        for run in hdr[j].paragraphs[0].runs:
            run.bold = True
    for c in component_results:
        row = table.add_row().cells
        row[0].text = c["domain"]
        row[1].text = c["workload"]
        row[2].text = f"{c['capacity']} {c.get('unit', '')}".strip()
        row[3].text = c.get("rationale", "")

    # 3. Executive summary table
    _h(doc, "3. Solution Summary")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for j, h in enumerate(["Domain", "Recommended Vendor", "3-Yr TCO (est.)", "Compliance"]):
        hdr[j].text = h
        for run in hdr[j].paragraphs[0].runs:
            run.bold = True
    stack_low = stack_high = 0
    for c in component_results:
        state = c["state"]
        top = _top_vendor(state)
        tco_by_vendor = {t["vendor"]: t for t in state.get("tco_estimates", [])}
        tco = tco_by_vendor.get(top.get("name")) if top else None
        row = table.add_row().cells
        row[0].text = c["domain"]
        row[1].text = top["name"] if top else "-"
        row[2].text = (f"{fmt_money(tco['range'][0])} – {fmt_money(tco['range'][1])}"
                       if tco else "-")
        vp = row[3].paragraphs[0]
        _status_run(vp, verdicts[c["domain"]])
        if tco:
            stack_low += tco["range"][0]
            stack_high += tco["range"][1]
    p = doc.add_paragraph()
    p.add_run("Combined stack 3-year TCO (top-choice vendors): ").bold = True
    p.add_run(f"{fmt_money(stack_low)} – {fmt_money(stack_high)}"
              if stack_low else fmt_money(synergy.get("estimated_stack_tco_3yr", 0)))
    p2 = doc.add_paragraph()
    p2.add_run("Consolidated compliance verdict: ").bold = True
    _status_run(p2, overall)
    if overall == "fail":
        p2.add_run("  (a solution is approvable only when every component passes)")

    # 4. Cross-domain vendor synergy
    _h(doc, "4. Cross-Domain Vendor Synergy")
    if synergy.get("multi_domain_vendors"):
        for vendor, hits in synergy["multi_domain_vendors"].items():
            p = doc.add_paragraph()
            p.add_run(vendor).bold = True
            p.add_run(" - recommended in: " + "; ".join(hits))
        for b in synergy.get("bundle_opportunities", []):
            doc.add_paragraph(b, style="List Bullet")
        for n in synergy.get("concentration_notes", []):
            cp = doc.add_paragraph(style="List Bullet")
            r = cp.add_run("Concentration: ")
            r.bold = True
            r.font.color.rgb = STATUS_COLORS["warn"]
            cp.add_run(n)
    else:
        doc.add_paragraph("No single vendor was recommended across multiple domains - "
                          "per-domain agreements apply.")
    _small_italic(doc,
        "Synergy analysis is deterministic (blueprints.analyze_synergy) - vendors are "
        "counted only where the recommendation matched the vendor registry and a computed "
        "TCO entry exists.")

    # 5+. Per-component detail sections
    section = 5
    for c in component_results:
        state = c["state"]
        arch = state.get("architecture_analysis", {})
        top = _top_vendor(state)
        tco_list = state.get("tco_estimates", [])

        _h(doc, f"{section}. {c['domain']} - {c['workload']}")

        if arch.get("architecture_type"):
            p = doc.add_paragraph()
            p.add_run(f"Architecture pattern: {arch['architecture_type']}").bold = True
        for rec in arch.get("key_recommendations", [])[:4]:
            doc.add_paragraph(rec, style="List Bullet")

        if top:
            p = doc.add_paragraph()
            p.add_run(f"Recommended option: {top['name']} "
                      f"(fit score {top.get('fit_score', 'N/A')}/10)").bold = True

        if tco_list:
            _h(doc, f"{section}.1 Three-Year TCO", level=2)
            _tco_table(doc, tco_list, fmt_money)

        _h(doc, f"{section}.2 Compliance", level=2)
        _compliance_block(doc, state.get("compliance_results", []))

        section += 1

    # Shared methodology note once, instead of per component
    _methodology_note(doc)

    # Next steps
    _h(doc, f"{section}. Recommendation and Next Steps")
    if overall == "fail":
        doc.add_paragraph(
            "One or more components currently fail compliance guardrails. Resolve the "
            "flagged items (or obtain documented ARB exceptions) before procurement.",
            style="List Number")
    if synergy.get("bundle_opportunities"):
        doc.add_paragraph(
            "Pursue bundled/stack-level commercial negotiation with multi-domain vendors "
            "identified in Section 4.", style="List Number")
    doc.add_paragraph("Engage shortlisted vendors for detailed sizing per component.",
                      style="List Number")
    doc.add_paragraph("Request formal quotations with 3-year TCO breakdown at solution level.",
                      style="List Number")
    doc.add_paragraph("Plan an integrated proof-of-concept exercising the cross-domain "
                      "sizing assumptions (driver metric at target value).",
                      style="List Number")

    _approvals(doc, ["Solution Architect", "Architecture Review Board Chair",
                     "Procurement / Vendor Management", "Business Sponsor"])
    _disclaimer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
